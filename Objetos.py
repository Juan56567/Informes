from docx import Document
from Enunciados import *


class Botella:
    def __init__(self, location):
        self.location = location
        self.doc = Document(self.location)
        self.tables = self.doc.tables
        self.camisa = self.tables[1].cell(26, 11)
        self.vastago = self.tables[1].cell(26, 25)
        self.piston = self.tables[2].cell(3, 0)
        self.montajeF = self.tables[2].cell(3, 1)
        self.montajeP = self.tables[2].cell(7, 0)
        self.tapa = self.tables[2].cell(7, 1)
        self.fecha = self.tables[1].cell(4, 25)
        self.componentes = [
            self.camisa,
            self.vastago,
            self.piston,
            self.montajeF,
            self.montajeP,
            self.tapa,
        ]

    def automatizar(self):
        ftemp = self.fecha.text.split(",")
        self.fecha.text = ftemp[0]

        for i in range(10, 18):
            read = self.tables[1].cell(i, 25)
            if len(read.text) < 2:
                read.text = "BUENO"

        # LEER CADA CELDA DE LA LISTA PREVIAMENTE CREADA
        for c in self.componentes:
            text = c.text.split("-")
            if len(text) > 1:
                if text[0] == "BUENO":
                    c.text = bueno
                elif text[0] == "CAMBIO":
                    if "CORROSION" in text[1]:
                        c.text = cambioCorrosion
                    elif "TOLERANCIA" in text[1]:
                        c.text = cambioTolerancia
                    elif "GOLPES" in text[1]:
                        c.text = cambioGolpes
                    elif "DESPRENDIMIENTO" in text[1]:
                        c.text = cambioDesprendimiento
                    elif "DESGASTE" in text[1]:
                        c.text = cambioDesgaste
                    elif "RALLAS" in text[1]:
                        c.text = cambioRallas
                    elif "FLEXION" in text[1]:
                        c.text = cambioFlexion
                    else:
                        pass
                else:
                    if "CROMAR" in text[1]:
                        c.text = rectiCromo
                    elif "BRILLAR" in text[1]:
                        c.text = rectiBrillar
                    elif "BRUÑIR" in text[1]:
                        c.text = rectiBrunir
                    elif "GOLPES" in text[1]:
                        c.text = rectiDeformacion
                    elif "ROSCA" in text[1]:
                        c.text = rectiRosca
            else:
                pass

        self.doc.save(self.location)

class Telescopico(Botella):
    def __init__(self, location):
        self.location = location
        self.doc = Document(self.location)
        self.tables = self.doc.tables
        self.camisa = self.tables[1].cell(27, 10)
        self.camisa1 = self.tables[1].cell(27, 18)
        self.camisa2 = self.tables[2].cell(3, 0)
        self.camisa3 = self.tables[2].cell(3, 1)
        self.montajeF = self.tables[2].cell(7, 0)
        self.montajeP = self.tables[2].cell(7, 1)
        self.tapa = self.tables[2].cell(11, 1)
        self.fecha = self.tables[1].cell(4, 18)
        self.componentes = [
            self.camisa,
            self.camisa1,
            self.camisa2,
            self.camisa3,
            self.montajeF,
            self.montajeP,
            self.tapa,
        ]

class Tirantes(Botella):
    def __init__(self, location):
        self.location = location
        self.doc = Document(self.location)
        self.tables = self.doc.tables
        self.camisa = self.tables[1].cell(27, 11)
        self.vastago = self.tables[1].cell(27, 22)
        self.piston = self.tables[2].cell(3, 0)
        self.montajeF = self.tables[2].cell(3, 1)
        self.montajeP = self.tables[2].cell(7, 0)
        self.tapa = self.tables[2].cell(7, 1)
        self.fecha = self.tables[1].cell(4, 25)
        self.componentes = [
            self.camisa,
            self.vastago,
            self.piston,
            self.montajeF,
            self.montajeP,
            self.tapa,
        ]

class Buzo(Botella):
    def __init__(self, location):
        self.location = location
        self.doc = Document(self.location)
        self.tables = self.doc.tables
        self.camisa = self.tables[1].cell(23, 10)
        self.vastago = self.tables[1].cell(23, 20)
        self.piston = self.tables[2].cell(3, 1)
        self.montajeF = self.tables[2].cell(3, 0)
        self.montajeP = self.tables[2].cell(7, 1)
        self.tapa = self.tables[2].cell(7, 0)
        self.fecha = self.tables[1].cell(4, 20)
        self.componentes = [
            self.camisa,
            self.vastago,
            self.piston,
            self.montajeF,
            self.montajeP,
            self.tapa,
        ]